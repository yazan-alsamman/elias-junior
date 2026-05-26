"""
Generate low-ATS-score sample CVs in cv-s/ (target app score ~5).

Score (Flutter): FAIL → max(5, 100 - 10*failed_rules - (25 if basic))

Run (ATS on :8000):
  cd cv-parser
  .venv\\Scripts\\activate
  pip install python-docx fpdf2 -q
  python ..\\cv-s\\generate_samples.py
"""
from __future__ import annotations

import json
import urllib.request
from pathlib import Path

OUT = Path(__file__).resolve().parent
_PAD_TO_BYTES = int(5.2 * 1024 * 1024)  # triggers E2_FILE_SIZE_LIMIT (>5 MB)


def _pad_pdf(path: Path) -> None:
    raw = path.read_bytes()
    if len(raw) < _PAD_TO_BYTES:
        path.write_bytes(raw + b"\n%" + b"0" * (_PAD_TO_BYTES - len(raw) - 2))


def make_pdf_score_05(path: Path) -> None:
    """7 failures → app score 5 (includes oversized file rule)."""
    from fpdf import FPDF

    pdf = FPDF()
    for _ in range(3):
        pdf.add_page()
        pdf.set_font("Helvetica", size=8)
        pdf.cell(0, 5, "x", new_x="LMARGIN", new_y="NEXT")
    pdf.output(str(path))
    _pad_pdf(path)


def make_pdf_score_15(path: Path) -> None:
    """6 failures, no padding → app score ~15."""
    from fpdf import FPDF

    pdf = FPDF()
    for _ in range(3):
        pdf.add_page()
        pdf.set_font("Helvetica", size=8)
        pdf.cell(0, 5, "x", new_x="LMARGIN", new_y="NEXT")
        for i in range(11):
            pdf.cell(0, 5, f"A{i}\tB{i}\tC{i}\tval", new_x="LMARGIN", new_y="NEXT")
    pdf.output(str(path))


def make_pdf_score_35(path: Path) -> None:
    """4 failures → app score ~35."""
    from fpdf import FPDF

    pdf = FPDF()
    for _ in range(3):
        pdf.add_page()
        pdf.set_font("Helvetica", size=9)
        for i in range(12):
            pdf.cell(
                0,
                5,
                f"Row{i}\tColA\tColB\tval",
                new_x="LMARGIN",
                new_y="NEXT",
            )
    pdf.output(str(path))


def make_docx_score_25(path: Path) -> None:
    """5 failures → app score ~25."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Taylor Kim \uFFFD summary")
    doc.add_paragraph("No standard Experience / Education / Skills headings.")
    t = doc.add_table(rows=4, cols=2)
    for r in range(4):
        for c in range(2):
            t.rows[r].cells[c].text = "data"
    doc.save(path)


def make_docx_score_55(path: Path) -> None:
    """2 failures → app score ~55."""
    from docx import Document

    doc = Document()
    doc.sections[0].header.paragraphs[0].text = (
        "Morgan Ellis | morgan.ellis@testmail.com | +1 555 442 1099"
    )
    doc.add_paragraph("Body without contact details or standard headings.")
    t = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            t.rows[r].cells[c].text = f"{r}-{c}"
    doc.save(path)


def ping_ats(file_path: Path, ats_base: str = "http://127.0.0.1:8000") -> dict | None:
    import mimetypes

    boundary = "----cvsgenboundary7MA4YWxk"
    data = file_path.read_bytes()
    mime = mimetypes.guess_type(file_path.name)[0] or "application/octet-stream"
    body = (
        f"--{boundary}\r\n"
        f'Content-Disposition: form-data; name="file"; filename="{file_path.name}"\r\n'
        f"Content-Type: {mime}\r\n\r\n"
    ).encode() + data + f"\r\n--{boundary}--\r\n".encode()
    req = urllib.request.Request(
        f"{ats_base.rstrip('/')}/ats-format/check",
        data=body,
        method="POST",
        headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            raw = json.loads(resp.read().decode())
    except Exception as err:
        print(f"  [ATS offline] {file_path.name}: {err}")
        return None
    n = int(raw.get("failed_rules_count") or 0)
    basic = bool(raw.get("failed_basic"))
    score = (
        100
        if raw.get("decision") == "PASS"
        else max(5, min(99, 100 - n * 10 - (25 if basic else 0)))
    )
    raw["_app_score"] = score
    return raw


def main() -> None:
    try:
        from docx import Document  # noqa: F401
        from fpdf import FPDF  # noqa: F401
    except ImportError:
        raise SystemExit("pip install python-docx fpdf2")

    samples = [
        ("cv_ats_score_05.pdf", make_pdf_score_05),
        ("cv_ats_score_05b.pdf", make_pdf_score_05),
        ("cv_ats_score_15.pdf", make_pdf_score_15),
        ("cv_ats_score_35.pdf", make_pdf_score_35),
        ("cv_ats_score_25.docx", make_docx_score_25),
        ("cv_ats_score_55.docx", make_docx_score_55),
    ]

    print(f"Output: {OUT}\n")
    rows = []
    for name, fn in samples:
        path = OUT / name
        fn(path)
        mb = path.stat().st_size / (1024 * 1024)
        print(f"  wrote {name} ({mb:.1f} MB)")
        r = ping_ats(path)
        if r:
            rows.append(
                {
                    "file": name,
                    "app_score": r["_app_score"],
                    "rules": r["failed_rules_count"],
                    "basic": r["failed_basic"],
                }
            )
            print(
                f"       -> app score {r['_app_score']} "
                f"({r['failed_rules_count']} rules)"
            )

    (OUT / "README.md").write_text(
        "\n".join(
            [
                "# cv-s — low ATS score test CVs",
                "",
                "Generated samples for upload testing. **Use `cv_ats_score_05.pdf` for score ~5.**",
                "",
                "| File | App score | Notes |",
                "|------|-----------|--------|",
                *[
                    f"| {x['file']} | **{x['app_score']}** | {x['rules']} failed rules |"
                    for x in rows
                ],
                "",
                "Score ~5 PDFs are ~5.2 MB (triggers file-size rule). Regenerate:",
                "",
                "```powershell",
                "cd cv-parser",
                ".venv\\Scripts\\activate",
                "pip install python-docx fpdf2 -q",
                "python ..\\cv-s\\generate_samples.py",
                "```",
                "",
                "Requires ATS: `http://127.0.0.1:8000` and app stack: `start-local-dev.cmd`",
                "",
            ]
        ),
        encoding="utf-8",
    )
    print(f"\n{OUT / 'README.md'}")


if __name__ == "__main__":
    main()
