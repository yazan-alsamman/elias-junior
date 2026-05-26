"""
ATS Rule Engine — Single-file FastAPI deployment.

Run on VPS:
    pip install fastapi uvicorn python-docx pypdf python-multipart pdfminer.six pydantic
    uvicorn ats_api:app --host 0.0.0.0 --port 8000

POST /ats-format/check  (multipart "file": PDF or DOCX)
GET  /ats-format/template
GET  /health
"""

from __future__ import annotations

import re
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Literal, Optional, Tuple

from docx import Document
from fastapi import FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pdfminer.high_level import extract_pages
from pdfminer.layout import LTFigure, LTImage, LTTextContainer, LTTextLine
from pydantic import BaseModel
from pypdf import PdfReader

# =====================================================================
#  Rules config (embedded — was ats_rules.json)
# =====================================================================

RULES: Dict[str, Any] = {
    "file": {
        "allowed_ext": [".pdf", ".docx"],
        "max_size_mb": 5,
        "max_pages_pdf": 2,
    },
    "pdf": {
        "scanned_min_chars_per_page": 180,
        "scanned_fail_ratio_pages_low_text": 0.6,
        "header_footer_band_ratio": 0.08,
        "max_images_any": 0,
        "columns": {
            "enabled": True,
            "min_lines_for_check": 35,
            "two_column_split_ratio": 0.16,
            "min_left_share": 0.30,
            "min_right_share": 0.30,
        },
        "tables": {
            "enabled": True,
            "min_table_like_lines": 10,
        },
    },
    "docx": {
        "max_tables_any": 0,
        "max_images_any": 0,
    },
    "contact": {
        "require_email": True,
        "require_phone": True,
        "email_regex": r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}",
        "phone_regex": r"(\+?\d{1,3}[\s.-]?)?(\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}",
    },
    "text_quality": {
        "max_replacement_char_ratio": 0.005,
    },
    "headings": {
        "required_any_of_groups": [
            ["experience", "work experience", "employment", "professional experience"],
            ["education", "academic"],
            ["skills", "technical skills", "core skills"],
        ],
    },
}

EMAIL_RE = re.compile(RULES["contact"]["email_regex"], re.IGNORECASE)
PHONE_RE = re.compile(RULES["contact"]["phone_regex"], re.IGNORECASE)

TEMPLATE_DIR = Path(__file__).parent / "ats_template_cache"
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATE_PATH = TEMPLATE_DIR / "ats_cv_template.docx"

# =====================================================================
#  Pydantic schemas
# =====================================================================


class FailureItem(BaseModel):
    rule_id: str
    severity: Literal["BASIC", "EXTRA"]
    issue: str
    fix: str


class Recommendation(BaseModel):
    type: Literal["template", "improve_current_cv", "none"]
    message: str
    improvements: List[str]


class ATSFormatResponse(BaseModel):
    filename: str
    decision: Literal["PASS", "FAIL"]
    failed_basic: bool
    failed_rules_count: int
    recommendation: Recommendation
    failures: List[FailureItem]
    template_url: Optional[str] = None


# =====================================================================
#  PDF / DOCX parser
# =====================================================================


@dataclass
class PDFLine:
    text: str
    x0: float
    x1: float
    y0: float
    y1: float


@dataclass
class FormatMetrics:
    file_type: str
    file_size_bytes: int

    pdf_pages: int = 0
    pdf_text_chars_per_page: Optional[List[int]] = None
    pdf_images_total: int = 0
    pdf_lines: Optional[List[Tuple[int, PDFLine]]] = None

    docx_tables: int = 0
    docx_images: int = 0
    docx_header_text: str = ""
    docx_footer_text: str = ""
    docx_body_text_sample: str = ""

    extracted_text_sample: str = ""
    extracted_text_full: str = ""


def _clean(s: str) -> str:
    return " ".join((s or "").replace("\u00a0", " ").split())


def _pdf_extract_layout(
    path: Path,
) -> Tuple[int, List[int], int, List[Tuple[int, PDFLine]], str, str]:
    reader = PdfReader(str(path))
    pages = len(reader.pages)

    chars_pp: List[int] = []
    sample_parts: List[str] = []
    full_parts: List[str] = []
    for page in reader.pages:
        t = _clean(page.extract_text() or "")
        chars_pp.append(len(t))
        if t:
            full_parts.append(t)
            if len(sample_parts) < 3:
                sample_parts.append(t[:900])

    lines: List[Tuple[int, PDFLine]] = []
    images_total = 0

    for page_index, layout in enumerate(extract_pages(str(path))):
        for element in layout:
            if isinstance(element, LTImage):
                images_total += 1
            if isinstance(element, LTFigure):
                for child in element:
                    if isinstance(child, LTImage):
                        images_total += 1

            if isinstance(element, LTTextContainer):
                for obj in element:
                    if isinstance(obj, LTTextLine):
                        text = _clean(obj.get_text())
                        if not text:
                            continue
                        lines.append(
                            (
                                page_index,
                                PDFLine(
                                    text=text,
                                    x0=float(obj.x0),
                                    x1=float(obj.x1),
                                    y0=float(obj.y0),
                                    y1=float(obj.y1),
                                ),
                            )
                        )

    sample = _clean(" ".join(sample_parts))[:2600]
    full_text = _clean("\n".join(full_parts))[:50000]
    return pages, chars_pp, images_total, lines, sample, full_text


def _docx_count_images(doc) -> int:
    try:
        return len(doc.inline_shapes)
    except Exception:
        return 0


def _docx_header_footer_text(doc) -> Tuple[str, str]:
    header_texts: List[str] = []
    footer_texts: List[str] = []
    try:
        for section in doc.sections:
            header_texts.extend([p.text for p in section.header.paragraphs])
            footer_texts.extend([p.text for p in section.footer.paragraphs])
    except Exception:
        pass
    return _clean(" ".join(header_texts)), _clean(" ".join(footer_texts))


def parse_format_metrics(file_path) -> FormatMetrics:
    path = Path(file_path)
    ext = path.suffix.lower()
    size = path.stat().st_size if path.exists() else 0

    if ext == ".pdf":
        pages, chars_pp, images_total, lines, sample, full_text = _pdf_extract_layout(
            path
        )
        return FormatMetrics(
            file_type="pdf",
            file_size_bytes=size,
            pdf_pages=pages,
            pdf_text_chars_per_page=chars_pp,
            pdf_images_total=images_total,
            pdf_lines=lines,
            extracted_text_sample=sample,
            extracted_text_full=full_text,
        )

    if ext == ".docx":
        doc = Document(str(path))
        header_text, footer_text = _docx_header_footer_text(doc)
        images = _docx_count_images(doc)
        body_text = _clean(" ".join(p.text for p in doc.paragraphs))
        sample = body_text[:2600]
        full_text = body_text[:50000]
        return FormatMetrics(
            file_type="docx",
            file_size_bytes=size,
            docx_tables=len(doc.tables),
            docx_images=images,
            docx_header_text=header_text,
            docx_footer_text=footer_text,
            docx_body_text_sample=sample,
            extracted_text_sample=sample,
            extracted_text_full=full_text,
        )

    return FormatMetrics(file_type="unknown", file_size_bytes=size)


# =====================================================================
#  Rule helpers
# =====================================================================


def _mb(size_bytes: int) -> float:
    return size_bytes / (1024 * 1024)


def _replacement_char_ratio(text: str) -> float:
    if not text:
        return 0.0
    bad = text.count("\uFFFD") + text.count("\ufffd")
    return bad / max(1, len(text))


def _has_required_headings(sample: str) -> Tuple[bool, List[str]]:
    s = (sample or "").lower()
    missing = []
    for group in RULES["headings"]["required_any_of_groups"]:
        if not any(h in s for h in group):
            missing.append(" / ".join(group))
    return (len(missing) == 0), missing


def _find_contacts(text: str) -> Dict[str, List[str]]:
    text = text or ""
    emails = list({m.group(0) for m in EMAIL_RE.finditer(text)})
    phones = list({m.group(0) for m in PHONE_RE.finditer(text)})
    return {"emails": emails, "phones": phones}


# ----- PDF checks -----


def _pdf_is_scanned(metrics: FormatMetrics) -> bool:
    chars_pp = metrics.pdf_text_chars_per_page or []
    if not chars_pp:
        return True
    min_chars = int(RULES["pdf"]["scanned_min_chars_per_page"])
    low_pages = sum(1 for c in chars_pp if c < min_chars)
    ratio = low_pages / max(1, metrics.pdf_pages)
    return ratio >= float(RULES["pdf"]["scanned_fail_ratio_pages_low_text"])


def _pdf_has_images(metrics: FormatMetrics) -> bool:
    return (metrics.pdf_images_total or 0) > int(RULES["pdf"]["max_images_any"])


def _pdf_columns_risk(metrics: FormatMetrics) -> bool:
    cfg = RULES["pdf"]["columns"]
    if not cfg.get("enabled", True):
        return False

    lines = metrics.pdf_lines or []
    if len(lines) < int(cfg["min_lines_for_check"]):
        return False

    by_page: Dict[int, List[PDFLine]] = {}
    for p, line in lines:
        by_page.setdefault(p, []).append(line)

    for _, page_lines in by_page.items():
        max_x1 = max((ln.x1 for ln in page_lines), default=0.0)
        if max_x1 <= 0:
            continue

        mid = max_x1 / 2.0
        split_tol = float(cfg["two_column_split_ratio"]) * max_x1

        left = right = total = 0
        for ln in page_lines:
            total += 1
            if ln.x0 < (mid - split_tol):
                left += 1
            elif ln.x0 > (mid + split_tol):
                right += 1

        if total == 0:
            continue

        left_share = left / total
        right_share = right / total

        if left_share >= float(cfg["min_left_share"]) and right_share >= float(
            cfg["min_right_share"]
        ):
            return True

    return False


def _pdf_contact_only_in_header_footer(metrics: FormatMetrics) -> bool:
    band_ratio = float(RULES["pdf"]["header_footer_band_ratio"])
    lines = metrics.pdf_lines or []
    if not lines:
        return False

    bounds: Dict[int, Tuple[float, float]] = {}
    for p, ln in lines:
        lo, hi = bounds.get(p, (1e9, -1e9))
        bounds[p] = (min(lo, ln.y0), max(hi, ln.y1))

    found_any = False
    found_in_body = False

    for p, ln in lines:
        page_lo, page_hi = bounds.get(p, (0.0, 1.0))
        height = max(1.0, page_hi - page_lo)
        top_band_start = page_hi - (band_ratio * height)
        bot_band_end = page_lo + (band_ratio * height)

        has_contact = bool(EMAIL_RE.search(ln.text) or PHONE_RE.search(ln.text))
        if not has_contact:
            continue

        found_any = True
        in_header = ln.y1 >= top_band_start
        in_footer = ln.y0 <= bot_band_end
        if not (in_header or in_footer):
            found_in_body = True

    return found_any and (not found_in_body)


def _pdf_table_risk(metrics: FormatMetrics) -> bool:
    cfg = RULES["pdf"]["tables"]
    if not cfg.get("enabled", True):
        return False

    lines = metrics.pdf_lines or []
    if not lines:
        return False

    table_like = 0
    for _, ln in lines:
        if "  " in ln.text or "\t" in ln.text:
            table_like += 1
        nums = re.findall(r"\d{2,4}", ln.text)
        if len(nums) >= 2:
            table_like += 1

    return table_like >= int(cfg["min_table_like_lines"])


# ----- DOCX checks -----


def _docx_has_images(metrics: FormatMetrics) -> bool:
    return (metrics.docx_images or 0) > int(RULES["docx"]["max_images_any"])


def _docx_has_tables(metrics: FormatMetrics) -> bool:
    return (metrics.docx_tables or 0) > int(RULES["docx"]["max_tables_any"])


def _docx_contact_only_in_header_footer(metrics: FormatMetrics) -> bool:
    header_contacts = _find_contacts(metrics.docx_header_text)
    footer_contacts = _find_contacts(metrics.docx_footer_text)
    body_contacts = _find_contacts(metrics.docx_body_text_sample)

    header_has = bool(header_contacts["emails"] or header_contacts["phones"])
    footer_has = bool(footer_contacts["emails"] or footer_contacts["phones"])
    body_has = bool(body_contacts["emails"] or body_contacts["phones"])

    return (header_has or footer_has) and (not body_has)


# =====================================================================
#  Main rule evaluator
# =====================================================================


def evaluate_ats_format(metrics: FormatMetrics) -> Dict[str, Any]:
    failures: List[Dict[str, str]] = []
    failed_basic = False

    def add_failure(rule_id: str, issue: str, fix: str, is_basic: bool):
        nonlocal failed_basic
        failures.append(
            {
                "rule_id": rule_id,
                "severity": "BASIC" if is_basic else "EXTRA",
                "issue": issue,
                "fix": fix,
            }
        )
        if is_basic:
            failed_basic = True

    if metrics.file_type not in ("pdf", "docx"):
        add_failure(
            "E1_ALLOWED_FILE_TYPE",
            "Unsupported file type.",
            "Convert the file to PDF or DOCX.",
            False,
        )
        failed_rules_count = len(failures)
        return {
            "decision": "FAIL",
            "failed_basic": False,
            "failed_rules_count": failed_rules_count,
            "recommendation": {
                "type": "template"
                if failed_rules_count > 5
                else "improve_current_cv",
                "message": "We recommend using an ATS-friendly template."
                if failed_rules_count > 5
                else "Apply the suggested fixes to your current CV.",
                "improvements": [f["fix"] for f in failures]
                if failed_rules_count <= 5
                else [],
            },
            "failures": failures,
        }

    max_size_mb = float(RULES["file"]["max_size_mb"])
    if _mb(metrics.file_size_bytes) > max_size_mb:
        add_failure(
            "E2_FILE_SIZE_LIMIT",
            "File size is too large.",
            f"Reduce the file size to less than {max_size_mb}MB.",
            False,
        )

    if metrics.file_type == "pdf":
        max_pages = int(RULES["file"]["max_pages_pdf"])
        if metrics.pdf_pages > max_pages:
            add_failure(
                "E3_MAX_PAGES_PDF",
                "PDF has too many pages.",
                f"Keep it to {max_pages} pages or fewer.",
                False,
            )

    if metrics.file_type == "pdf" and _pdf_has_images(metrics):
        add_failure(
            "B1_NO_IMAGES",
            "PDF contains images/icons/graphics.",
            "Remove images and icons.",
            True,
        )
    if metrics.file_type == "docx" and _docx_has_images(metrics):
        add_failure(
            "B1_NO_IMAGES",
            "DOCX contains images/icons.",
            "Remove images and icons.",
            True,
        )

    if metrics.file_type == "pdf" and _pdf_is_scanned(metrics):
        add_failure(
            "B6_NOT_SCANNED_PDF",
            "PDF appears scanned or text is not extractable.",
            "Export a text-based PDF from Word/Docs.",
            True,
        )

    if metrics.file_type == "pdf" and _pdf_columns_risk(metrics):
        add_failure(
            "B2_NO_COLUMNS",
            "High risk of multi-column layout.",
            "Use a single-column layout.",
            True,
        )

    if metrics.file_type == "pdf" and _pdf_contact_only_in_header_footer(metrics):
        add_failure(
            "B3_CONTACT_NOT_IN_HEADER_FOOTER_ONLY",
            "Contact info appears only in header/footer.",
            "Place contact info in the main body.",
            True,
        )
    if metrics.file_type == "docx" and _docx_contact_only_in_header_footer(metrics):
        add_failure(
            "B3_CONTACT_NOT_IN_HEADER_FOOTER_ONLY",
            "Contact info appears only in header/footer.",
            "Place contact info in the main body.",
            True,
        )

    if metrics.file_type == "docx" and _docx_has_tables(metrics):
        add_failure(
            "B4_NO_TABLES",
            "DOCX contains tables.",
            "Remove tables and use bullet points.",
            True,
        )
    if metrics.file_type == "pdf" and _pdf_table_risk(metrics):
        add_failure(
            "B4_NO_TABLES",
            "PDF shows table-like content.",
            "Avoid tables.",
            True,
        )

    if (
        metrics.file_type == "pdf"
        and (_pdf_has_images(metrics) or _pdf_is_scanned(metrics))
    ) or (metrics.file_type == "docx" and _docx_has_images(metrics)):
        add_failure(
            "B5_NO_TEXT_IN_IMAGES_ICONS",
            "Some content may be inside images/icons.",
            "Ensure all important content is real selectable text.",
            True,
        )

    contact_text = metrics.extracted_text_sample
    if metrics.file_type == "docx":
        contact_text = " ".join(
            [
                metrics.docx_header_text,
                metrics.docx_footer_text,
                metrics.docx_body_text_sample,
            ]
        )
    contacts = _find_contacts(contact_text)

    if RULES["contact"]["require_email"] and not contacts["emails"]:
        add_failure(
            "E4_CONTACT_EXISTS",
            "No clear email found.",
            "Add an email as plain text in the body.",
            False,
        )
    if RULES["contact"]["require_phone"] and not contacts["phones"]:
        add_failure(
            "E4_CONTACT_EXISTS",
            "No clear phone number found.",
            "Add a phone number as plain text in the body.",
            False,
        )

    ratio = _replacement_char_ratio(metrics.extracted_text_sample)
    if ratio > float(RULES["text_quality"]["max_replacement_char_ratio"]):
        add_failure(
            "E5_TEXT_ENCODING_OK",
            "Encoding/font issue detected (replacement characters).",
            "Use standard fonts like Arial/Calibri and export a text-based PDF.",
            False,
        )

    ok_head, missing = _has_required_headings(metrics.extracted_text_sample)
    if not ok_head:
        add_failure(
            "E6_HEADINGS_PRESENT",
            "Standard section headings are missing.",
            f"Add headings such as: {', '.join(missing)}",
            False,
        )

    if failures:
        failed_rules_count = len(failures)
        unique_improvements = list(dict.fromkeys(f["fix"] for f in failures))
        return {
            "decision": "FAIL",
            "failed_basic": failed_basic,
            "failed_rules_count": failed_rules_count,
            "recommendation": {
                "type": "template"
                if failed_rules_count > 5
                else "improve_current_cv",
                "message": "We recommend using an ATS-friendly template."
                if failed_rules_count > 5
                else "Improve your current CV using the following recommendations.",
                "improvements": unique_improvements
                if failed_rules_count <= 5
                else unique_improvements[:5],
            },
            "failures": failures,
        }

    return {
        "decision": "PASS",
        "failed_basic": False,
        "failed_rules_count": 0,
        "recommendation": {
            "type": "none",
            "message": "Your CV is compatible with the current ATS rules.",
            "improvements": [],
        },
        "failures": [],
    }


# =====================================================================
#  Service — accept upload, run checks, return payload
# =====================================================================


def ensure_template_exists() -> None:
    if TEMPLATE_PATH.exists():
        return

    doc = Document()
    doc.add_heading("ATS-Friendly CV Template", level=1)

    doc.add_paragraph("Name: ____________________________")
    doc.add_paragraph("Email: ____________________________")
    doc.add_paragraph("Phone: ____________________________")
    doc.add_paragraph("Location: ____________________________")
    doc.add_paragraph("LinkedIn/GitHub (optional): ____________________________")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "- 2-3 lines summary (plain text). No icons, no tables, no columns."
    )

    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Company - Role | City, Country | YYYY-YYYY")
    doc.add_paragraph("- Achievement bullet")
    doc.add_paragraph("- Achievement bullet")

    doc.add_heading("Education", level=2)
    doc.add_paragraph("University - Degree | YYYY-YYYY")

    doc.add_heading("Skills", level=2)
    doc.add_paragraph("- Skill 1, Skill 2, Skill 3, ...")

    doc.add_heading("Projects", level=2)
    doc.add_paragraph("Project Name - Role/Tech Stack | YYYY")
    doc.add_paragraph("- Brief impact-focused bullet.")

    doc.add_heading("Certifications", level=2)
    doc.add_paragraph("Certification Name - Issuer - Year")

    doc.save(str(TEMPLATE_PATH))


def process_cv_file(upload_file, base_url: str) -> Dict[str, Any]:
    ensure_template_exists()

    filename = (upload_file.filename or "").lower()
    suffix = Path(filename).suffix

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp_path = tmp.name

    try:
        with open(tmp_path, "wb") as out:
            shutil.copyfileobj(upload_file.file, out)

        metrics = parse_format_metrics(tmp_path)
        result = evaluate_ats_format(metrics)

        extracted_text = (metrics.extracted_text_full or metrics.extracted_text_sample or "").strip()
        payload = {
            "filename": upload_file.filename,
            "decision": result["decision"],
            "failed_basic": result["failed_basic"],
            "failed_rules_count": result["failed_rules_count"],
            "recommendation": result["recommendation"],
            "failures": result["failures"],
            "extracted_text": extracted_text[:50000] if extracted_text else "",
        }

        if result["decision"] == "FAIL" and result["failed_rules_count"] > 5:
            payload["template_url"] = (
                base_url.rstrip("/") + "/ats-format/template"
            )

        return payload

    finally:
        try:
            Path(tmp_path).unlink(missing_ok=True)
        except Exception:
            pass


# =====================================================================
#  FastAPI app
# =====================================================================

app = FastAPI(title="ATS Rule Engine", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
def root():
    return {"service": "ATS Rule Engine", "status": "ok"}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/ats-format/check", response_model=ATSFormatResponse)
async def check_ats_format(request: Request, file: UploadFile = File(...)):
    """
    Upload CV (PDF/DOCX) -> run ATS format checks -> return PASS/FAIL + issues.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")
    base_url = str(request.base_url)
    return process_cv_file(file, base_url)


@app.get("/ats-format/template")
def download_template():
    """
    Download ATS-friendly DOCX template.
    """
    ensure_template_exists()
    return FileResponse(
        path=str(TEMPLATE_PATH),
        filename="ats_cv_template.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
